"""
Document Parser Module

Parses PDF, DOCX, and Markdown documents to extract text content.
Uses factory pattern for extensibility.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup
import re
import logging

logger = logging.getLogger(__name__)


class DocumentParser(ABC):
    """Abstract base class for document parsers."""

    @abstractmethod
    def parse(self, file_path: str) -> str:
        """
        Parse document and return extracted text.

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text content
        """
        pass

    @abstractmethod
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract document metadata.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary of metadata
        """
        pass


class PDFParser(DocumentParser):
    """Parser for PDF documents using PyMuPDF."""

    def parse(self, file_path: str) -> str:
        """Extract text from PDF, preserving page structure."""
        doc = fitz.open(file_path)
        text_parts = []

        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text}")

        doc.close()
        return "\n\n".join(text_parts)

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF metadata."""
        doc = fitz.open(file_path)
        metadata = {
            "page_count": doc.page_count,
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "creator": doc.metadata.get("creator", ""),
            "producer": doc.metadata.get("producer", ""),
        }
        doc.close()
        return metadata


class DocxParser(DocumentParser):
    """Parser for Microsoft Word documents."""

    def parse(self, file_path: str) -> str:
        """Extract text from DOCX, preserving heading structure."""
        doc = DocxDocument(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                # Preserve heading information with markdown-style formatting
                if para.style and para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level_num = int(level)
                        prefix = '#' * min(level_num, 6)
                        text_parts.append(f"\n{prefix} {para.text}\n")
                    except ValueError:
                        text_parts.append(f"\n## {para.text}\n")
                else:
                    text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                text_parts.append("\n" + "\n".join(table_text) + "\n")

        return "\n".join(text_parts)

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DOCX metadata."""
        doc = DocxDocument(file_path)
        core_props = doc.core_properties

        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }


class MarkdownParser(DocumentParser):
    """Parser for Markdown documents."""

    def parse(self, file_path: str) -> str:
        """
        Extract text from Markdown.

        Converts to HTML first to handle formatting, then extracts plain text
        while preserving structure.
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convert to HTML then extract text to preserve structure
        html = markdown.markdown(
            md_content,
            extensions=['tables', 'fenced_code', 'nl2br']
        )
        soup = BeautifulSoup(html, 'html.parser')

        # Get text with newlines between blocks
        return soup.get_text(separator='\n')

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract Markdown metadata."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Extract title from first H1
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)

        # Count headings at each level
        h1_count = len(re.findall(r'^#\s', content, re.MULTILINE))
        h2_count = len(re.findall(r'^##\s', content, re.MULTILINE))
        h3_count = len(re.findall(r'^###\s', content, re.MULTILINE))

        # Count code blocks
        code_block_count = len(re.findall(r'```', content)) // 2

        return {
            "title": title_match.group(1) if title_match else "",
            "line_count": len(content.splitlines()),
            "character_count": len(content),
            "h1_count": h1_count,
            "h2_count": h2_count,
            "h3_count": h3_count,
            "code_block_count": code_block_count,
        }


class TextParser(DocumentParser):
    """Parser for plain text documents."""

    def parse(self, file_path: str) -> str:
        """Read plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract text file metadata."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return {
            "line_count": len(content.splitlines()),
            "character_count": len(content),
            "word_count": len(content.split()),
        }


class ParserFactory:
    """Factory for creating document parsers based on file type."""

    _parsers = {
        '.pdf': PDFParser,
        '.docx': DocxParser,
        '.doc': DocxParser,  # May not work for older .doc files
        '.md': MarkdownParser,
        '.markdown': MarkdownParser,
        '.txt': TextParser,
    }

    @classmethod
    def get_parser(cls, file_path: str) -> DocumentParser:
        """
        Get appropriate parser for file type.

        Args:
            file_path: Path to the document file

        Returns:
            DocumentParser instance

        Raises:
            ValueError: If file type is not supported
        """
        ext = Path(file_path).suffix.lower()
        parser_class = cls._parsers.get(ext)

        if not parser_class:
            raise ValueError(f"Unsupported file type: {ext}. Supported types: {cls.supported_extensions()}")

        return parser_class()

    @classmethod
    def supported_extensions(cls) -> list:
        """Return list of supported file extensions."""
        return list(cls._parsers.keys())

    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if a file type is supported."""
        ext = Path(file_path).suffix.lower()
        return ext in cls._parsers


def parse_document(file_path: str) -> tuple[str, Dict[str, Any]]:
    """
    Convenience function to parse a document and extract metadata.

    Args:
        file_path: Path to the document file

    Returns:
        Tuple of (text_content, metadata_dict)
    """
    parser = ParserFactory.get_parser(file_path)
    text = parser.parse(file_path)
    metadata = parser.extract_metadata(file_path)
    return text, metadata
